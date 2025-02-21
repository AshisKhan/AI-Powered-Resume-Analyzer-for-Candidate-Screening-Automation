<<<<<<< HEAD
import os
import textract
import fitz
from docx import Document
import re



def clean_text(text):
    # Remove newlines, tabs, and extra spaces
    text = text.replace('\n', ' ').replace('\r', '').replace('\t', ' ')
    
    # Remove any non-alphanumeric characters (except spaces and common punctuation)
    text = re.sub(r'[^A-Za-z0-9\s.,!?\'"-@]', '', text)
    
    # Remove multiple spaces between words
    text = re.sub(r'\s+', ' ', text)
    
    # Trim leading and trailing spaces
    text = text.strip()
    # Remove common headers/footers (if any identified patterns like 'Confidential', 'Resume', etc.)
    text = re.sub(r'\b(Confidential|Resume|Curriculum Vitae)\b', '', text, flags=re.IGNORECASE)
   
    
    return text.lower()



def read_file(file_path):
    file_type = os.path.splitext(file_path)

    # PDF file extraction
    if file_type[1].lower() == ".pdf":
        try:
            # Open text-based PDFs using PyMuPDF
            data = fitz.open(file_path)
            text = ''
            for page in data:
                text += page.get_text()
                text = re.sub(r'Page \d+', '', text)
            return text
        
        except Exception as e:
            print(f"PyMuPDF failed: {e}. Trying Textract...")
            try:
                # Fallback to Textract for scanned PDFs
                return textract.process(file_path).decode("utf-8")
            except Exception as ex:
                print(f"Textract failed: {ex}")
                return None

    # DOCX file extraction
    elif file_type[1].lower() == ".docx":
        try:
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return None

    # TXT file extraction
    elif file_type[1].lower() == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return None

    # Fallback for unsupported formats
    else:
        try:
            return textract.process(file_path).decode("utf-8")
        except Exception as e:
            print(f"Error processing file: {e}")
            return None





=======
import os
import textract
import fitz
from docx import Document
import re



def clean_text(text):
    # Remove newlines, tabs, and extra spaces
    text = text.replace('\n', ' ').replace('\r', '').replace('\t', ' ')
    
    # Remove any non-alphanumeric characters (except spaces and common punctuation)
    text = re.sub(r'[^A-Za-z0-9\s.,!?\'"-@]', '', text)
    
    # Remove multiple spaces between words
    text = re.sub(r'\s+', ' ', text)
    
    # Trim leading and trailing spaces
    text = text.strip()
    # Remove common headers/footers (if any identified patterns like 'Confidential', 'Resume', etc.)
    text = re.sub(r'\b(Confidential|Resume|Curriculum Vitae)\b', '', text, flags=re.IGNORECASE)
   
    
    return text.lower()



def read_file(file_path):
    file_type = os.path.splitext(file_path)

    # PDF file extraction
    if file_type[1].lower() == ".pdf":
        try:
            # Open text-based PDFs using PyMuPDF
            data = fitz.open(file_path)
            text = ''
            for page in data:
                text += page.get_text()
                text = re.sub(r'Page \d+', '', text)
            return text
        
        except Exception as e:
            print(f"PyMuPDF failed: {e}. Trying Textract...")
            try:
                # Fallback to Textract for scanned PDFs
                return textract.process(file_path).decode("utf-8")
            except Exception as ex:
                print(f"Textract failed: {ex}")
                return None

    # DOCX file extraction
    elif file_type[1].lower() == ".docx":
        try:
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return None

    # TXT file extraction
    elif file_type[1].lower() == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return None

    # Fallback for unsupported formats
    else:
        try:
            return textract.process(file_path).decode("utf-8")
        except Exception as e:
            print(f"Error processing file: {e}")
            return None





>>>>>>> 9b81f7a7d3c44723808206d97a1f66adffe7e7eb
